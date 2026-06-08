#!/usr/bin/env python3
"""Конвертер записки «Свод Сводов» из Markdown в .docx.

Поддерживает ровно те конструкции, что нужны для служебной записки:
  [RIGHT]текст     — абзац по правому краю
  [CENTER]текст    — абзац по центру
  **жирный**       — жирный фрагмент; если это число со знаком,
                     красится: «+...» зелёным, «-...» красным
  | a | b | c |     — таблица (со строкой-разделителем `:-:`)

Запуск:  python md_to_docx.py input.md "Свод Сводов за Май 2026.docx"
"""
from __future__ import annotations

import re
import sys

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

GREEN = RGBColor(0x00, 0x99, 0x00)
RED = RGBColor(0xC0, 0x00, 0x00)

SIGNED_NUM = re.compile(r"^[+\-−][\s \d.,%]*\d")


def _color_for(text: str):
    """Цвет для жирного фрагмента-числа: + зелёный, - красный, иначе None."""
    t = text.strip()
    if not SIGNED_NUM.match(t):
        return None
    return GREEN if t[0] == "+" else RED


def add_runs(paragraph, text: str, base_bold: bool = False) -> None:
    """Разбивает строку по ** и добавляет ранами с жирностью и цветом."""
    for i, chunk in enumerate(text.split("**")):
        if chunk == "":
            continue
        bold = base_bold or (i % 2 == 1)  # нечётные сегменты — внутри **...**
        run = paragraph.add_run(chunk)
        run.bold = bold
        if bold:
            color = _color_for(chunk)
            if color is not None:
                run.font.color.rgb = color


def is_table_line(line: str) -> bool:
    return line.lstrip().startswith("|")


def is_separator(line: str) -> bool:
    return bool(re.match(r"^\s*\|?[\s:\-|]+\|?\s*$", line)) and "-" in line


def split_cells(line: str):
    s = line.strip()
    if s.startswith("|"):
        s = s[1:]
    if s.endswith("|"):
        s = s[:-1]
    return [c.strip() for c in s.split("|")]


def add_table(doc: Document, rows: list[str]) -> None:
    parsed = [split_cells(r) for r in rows if not is_separator(r)]
    ncols = max(len(r) for r in parsed)
    table = doc.add_table(rows=0, cols=ncols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ri, cells in enumerate(parsed):
        row = table.add_row()
        for ci in range(ncols):
            text = cells[ci] if ci < len(cells) else ""
            cell = row.cells[ci]
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_runs(para, text, base_bold=(ri == 0))


def convert(md_path: str, docx_path: str) -> None:
    with open(md_path, encoding="utf-8") as f:
        lines = f.read().splitlines()

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        if is_table_line(line):
            block = []
            while i < n and is_table_line(lines[i]):
                block.append(lines[i])
                i += 1
            add_table(doc, block)
            doc.add_paragraph("")
            continue

        stripped = line.strip()
        if stripped == "":
            i += 1
            continue

        para = doc.add_paragraph()
        if stripped.startswith("[RIGHT]"):
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            add_runs(para, stripped[len("[RIGHT]"):])
        elif stripped.startswith("[CENTER]"):
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_runs(para, stripped[len("[CENTER]"):])
        else:
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            add_runs(para, stripped)
        i += 1

    doc.save(docx_path)
    print(f"OK: {docx_path}")


if __name__ == "__main__":
    src = sys.argv[1] if len(sys.argv) > 1 else "report.md"
    dst = sys.argv[2] if len(sys.argv) > 2 else "report.docx"
    convert(src, dst)
